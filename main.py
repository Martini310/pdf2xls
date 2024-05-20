import os
import re
import json
from datetime import date
from typing import List, Dict, Pattern
import argparse
from PIL import Image
import pdf2image
import dotenv
from openpyxl import Workbook
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import pytesseract
import PyPDF2


dotenv.load_dotenv()

pytesseract.pytesseract.tesseract_cmd = os.environ.get('TESSERACT_CMD')

CUSTOM_CONFIG = r'--oem 3 --psm 6 -l pol'
poppler_path=os.environ.get('POPPLER_PATH')


class PDFHandler:
    """
    A class to read text from a PDF file, extract certain data patterns, and create a docx file.
    
    Args:
        path (str): The path to the PDF file.
        scan (bool, optional): Whether to perform OCR on scanned PDFs. 
                            Read text from PDF if false. Defaults to False.
    """
    name_ptrn = r'A-Za-zĄąĆćĘęŁłŃńÓóŚśŹźŻż”\"\'©—&-'
    patterns: Dict[str, Pattern[str]] = {
        'kt': r'(?<=5410.[0-9].)\s*([0-9]+)',
        'name': rf'(?<=na rzecz )([{name_ptrn}]+(?:\s+[{name_ptrn}]+)*(?:\s+[{name_ptrn}]+))',
        'vin': r'(?<=V[I!]N[:;])[\s —-]*([\w?—-]*)',
        'basis': r'(?<=w związku z art\. )[\w\s\.]*(?= ustawy)',
        'tr': r'(?<=rej\.)\s*([A-Z0-9]+\s*[A-Z0-9]+)\b',
        'address': r'(?<=na rzecz )[\w\W]*(?=w związku)',
        'date': r'(?<=Poznań, dnia ).+(?=r)',
        'brand': r'(?<=marki\s)[\w\s\\/-—]+(?=o)',
        'pesel': r'[0-9]{9,11}',
        'purchase_date': r'(?<=z[\s]dnia)\s*[0-9-—/.\s]+(?=r.)',
    }

    def __init__(self, path: str, scan: bool = False) -> None:
        self.path: str = path
        self.scan: bool = scan
        self.text = self.extract_text() if not self.scan else self.perform_ocr(self.create_images())
        self.results = self.extract_data()
        self.process_results()

    def extract_text(self) -> str:
        """Extract text from PDF file"""
        with open(self.path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            return ''.join(page.extract_text() for page in pdf_reader.pages)

    def create_images(self) -> List[Image.Image]:
        """Convert PDF (scanned file) file into images."""
        return pdf2image.convert_from_path(self.path, poppler_path=poppler_path)

    def perform_ocr(self, images: List[Image.Image]) -> str:
        """Perform OCR on images."""
        return '\n'.join(pytesseract.image_to_string(page, config=CUSTOM_CONFIG) for page in images)

    def find_pattern(self, text: str, pattern: Pattern[str]) -> str:
        """Find and return the provided pattern in the given text."""
        matches: List[str] = re.findall(pattern, text)
        return matches[0].strip().strip('.').strip(',').replace('\n', ' ') if matches else 'null'

    def extract_data(self) -> Dict[str, str]:
        """Extract data from the text based on predefined patterns."""
        return {key: self.find_pattern(self.text, pattern) for key,pattern in self.patterns.items()}

    def process_results(self) -> None:
        """Process and format extracted results."""
        self.assign_activity()
        self.format_kt()
        self.format_date('date')
        self.format_date('purchase_date')
        self.format_vin()

    def assign_activity(self) -> None:
        """Assign an activity based on the legal basis."""
        basis = self.results['basis']
        if '73aa ust. 1 pkt 3' in basis:
            self.results['czynność'] = 'SPROWADZONY'
        elif '73aa ust. 1 pkt 1' in basis:
            self.results['czynność'] = 'NABYCIE'
        elif '78 ust. 2 pkt 1' in basis:
            self.results['czynność'] = 'ZBYCIE'
        else:
            self.results['czynność'] = 'null'

    def format_kt(self) -> None:
        """Format the KT number to be five digits long."""
        kt = self.results['kt']
        self.results['kt'] = kt.zfill(5) if kt != 'null' and len(kt) < 5 else kt

    def format_date(self, name: str) -> None:
        """Replace wrong characters in date and update the result."""
        dt = self.results[name]
        if dt != 'null':
            self.results[name] = dt.replace('—', '.').replace('-', '.').replace('/', '.')

    def format_vin(self) -> None:
        """Correct common mistakes in VIN pattern."""
        vin = self.results['vin']
        if vin != 'null':
            self.results['vin'] = vin.replace('O', '0')

    def initialize_doc(self) -> Document:
        """Initialize the Word document with custom styles and layout."""
        doc = Document()
        styles = doc.styles
        self.setup_title_style(styles)
        self.setup_normal_style(styles)
        self.setup_page_layout(doc.sections[0])
        return doc

    def setup_title_style(self, styles) -> None:
        """Set up the title style."""
        title_style = styles.add_style('Tytuł', WD_STYLE_TYPE.PARAGRAPH)
        title_style.base_style = styles['Normal']
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.font.bold = True

    def setup_normal_style(self, styles) -> None:
        """Set up the normal style."""
        normal_style = styles['Normal']
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal_style.paragraph_format.space_after = Cm(0)
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(10)

    def setup_page_layout(self, section) -> None:
        """Set up the page layout and margins."""
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def add_content(self, doc: Document, data: Dict[str, str], source: Dict[str, str]) -> None:
        """Add content to the Word document based on the extracted data."""
        # Choose right template
        if self.results['czynność'] == 'NABYCIE':
            kara = source['kara_nabycie']
            uzasadnienie = source['uzasadnienie_nabycie']
            uzasadnienie += source['uzasadnienie_wspolne']
        elif self.results['czynność'] == 'SPROWADZONY':
            kara = source['kara_ue']
            uzasadnienie = source['uzasadnienie_ue']
            uzasadnienie += source['uzasadnienie_wspolne']
        elif self.results['czynność'] == 'ZBYCIE':
            kara = source['kara_zbycie']
            uzasadnienie = source['uzasadnienie_zbycie']
            uzasadnienie += source['uzasadnienie_wspolne'][3:]
            source['podstawa_prawna'] = source['podstawa_prawna'].replace('140mb ust. 1', '140mb ust. 6')

        today = date.today()
        formatted_date = f"Poznań dnia {today.strftime('%d.%m.%Y')}r."

        doc.add_paragraph(formatted_date).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        header = 'Starosta Poznański\nul. Jackowskiego 18\n60-509 Poznań'
        doc.add_paragraph(header).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph(f"\n\nDECYZJA NR KT.5410.7.{data['kt']}.2024\n", style='Tytuł')

        doc.add_paragraph(source['podstawa_prawna'].format(
            data['basis'], data['name'], data['pesel'], data['brand'], data['tr'], data['vin']
            ))

        doc.add_paragraph('\nStarosta\n', style='Tytuł')
        doc.add_paragraph(kara)
        doc.add_paragraph('\nUzasadnienie\n', style='Tytuł')

        doc.add_paragraph(uzasadnienie[0].format(data['name'], data['purchase_date']))
        doc.add_paragraph(uzasadnienie[1])
        doc.add_paragraph(uzasadnienie[2].format(data['date']))
        for uzasadnienie in uzasadnienie[3:]:
            doc.add_paragraph(uzasadnienie)

        add_numbered_paragraphs(doc, source['przepisy'][:7], 'ListNumber', Inches(0.5))

        for przepis in source['przepisy'][7:10]:
            doc.add_paragraph(przepis)

        add_numbered_paragraphs(doc, source['przepisy'][10:12], 'List Number 2', space_after=Cm(0))
        paragraph = doc.add_paragraph(source['przepisy'][12])

        for przepis in source['przepisy'][13:15]:
            paragraph = doc.add_paragraph(przepis)
            paragraph.paragraph_format.left_indent = Inches(0.25)

        for przepis in source['przepisy'][15:]:
            paragraph = doc.add_paragraph(przepis)

        doc.add_paragraph('\nPouczenie\n', style='Tytuł')
        doc.add_paragraph(source['pouczenia'][0])
        paragraph = doc.add_paragraph('\n')

        paragraph.add_run('\tWpłaty należy dokonać na konto numer: ')
        paragraph.add_run('7710 3012 4700 0000 0034 9162 41').bold = True
        paragraph.add_run(' w tytule podając nr decyzji ')
        paragraph.add_run(f"KT.5410.7.{data['kt']}.2024").bold = True
        paragraph = doc.add_paragraph()

        doc.add_paragraph(source['pouczenia'][2])
        doc.add_paragraph()
        doc.add_paragraph(source['pouczenia'][3])
        doc.add_paragraph('\n' * 9)

        doc.add_paragraph('Otrzymują:')

        receivers = [data['address'], 'WYDZIAŁ FINANSÓW W MIEJSCU', 'a/a']
        add_numbered_paragraphs(doc, receivers, 'List Number 3', Inches(0.5))

        doc.add_paragraph('\nSprawę prowadzi:   Beata Andrzejewska tel. 61 8410 568')

    def save_doc(self, doc: Document, kt: str) -> None:
        """Save the Word document with a filename based on the KT number."""
        doc.save(f"{kt}_postanowienie.docx")

    def create_docx(self) -> None:
        '''Create an administrative decision imposing a penalty in .docx format'''
        try:
            is_valid(self)

            data = self.results
            with open('docx_source_text.json', 'r', encoding='utf-8') as file:
                source = json.load(file)
            doc = self.initialize_doc()
            self.add_content(doc, data, source)

            if data['kt'] == 'null':
                file_name = f"KT.5410.7.{data['tr']}.2024.docx"
            else:
                file_name = f"KT.5410.7.{data['kt']}.2024.docx"

            doc.save(file_name)
        except ValueError as e:
            print(e)


    def __str__(self) -> str:
        return '\n'.join(f'{key} - {value}' for key,value in self.results.items()) + '\n' + '-' * 50


class ReadPDF:
    """
    A class to read multiple PDF files, extract data using PDFHandler, and write to Excel.

    Args:
        path (str): The path to the directory containing PDF files.
        scan (bool, optional): Whether to perform OCR on scanned PDFs. Defaults to False.
        reverse (bool, optional): Whether to reverse the order of extracted data. Defaults to False.
    """
    def __init__(self, path: str, scan: bool = False, reverse: bool = False) -> None:
        self.path: str = path
        self.reverse: bool = reverse
        self.scan: bool = scan
        self.handlers: List[PDFHandler] = []
        self.files_paths: List[str] = [
            os.path.join(self.path, f)
            for f in os.listdir(self.path)
            if os.path.isfile(os.path.join(self.path, f)) and f.endswith('.pdf')
        ]

    def read_pdf(self) -> None:
        '''Create a list of PDFHandler objects and store it in self.handlers'''
        self.handlers = [PDFHandler(file_path, scan=self.scan) for file_path in self.files_paths]
        if self.reverse:
            self.handlers.reverse()

    def write_to_excel(self, data: List[PDFHandler], excel_file_path: str) -> None:
        """
        Write data extracted from OCR to an Excel file.

        Args:
            data (List[PDFHandler]): List of PDFHandler objects containing extracted data.
            excel_file_path (str): The path to save the Excel file.
        """
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = 'Extracted data (OCR)'

        row = 1
        for handler in data:
            try:
                # Check if a file is valid pdf
                is_valid(handler)

                kt = handler.results
                sheet[f'A{row}'] = kt['kt']
                sheet[f'B{row}'] = kt['tr']
                sheet[f'C{row}'] = kt['vin']
                sheet[f'D{row}'] = kt['name']
                sheet[f'E{row}'] = ''
                sheet[f'F{row}'] = ''
                sheet[f'G{row}'] = kt['date']
                sheet[f'H{row}'] = ''
                sheet[f'I{row}'] = ''
                sheet[f'J{row}'] = ''
                sheet[f'K{row}'] = kt['czynność']
                sheet[f'L{row}'] = kt['basis']

                row += 1
            except ValueError as e:
                print(e)
                continue
        workbook.save(excel_file_path)

def is_valid(pdf_obj: PDFHandler) -> None:
    """
    Check if all values in PDFHandler.results are 'null'. 
    If yes it's probably not a valid pdf object
    """
    if all(value == 'null' for value in pdf_obj.results.values()):
        raise ValueError(
            f"Prawdopodobnie plik '{pdf_obj.path}' nie "
            f"jest postanowieniem o wszczęciu postępowania.")
    return True

def add_numbered_paragraphs(doc, items, style_name, left_indent=None, space_after=None):
    """
    Add numbered paragraphs to the document with specified style and formatting.
    """
    for item in items:
        paragraph = doc.add_paragraph(item, style=style_name)
        if left_indent is not None:
            paragraph.paragraph_format.left_indent = left_indent
        if space_after is not None:
            paragraph.paragraph_format.space_after = space_after


if __name__ == '__main__':
    # test = PDFHandler('./skany/skany_OCR/20240515115607.pdf')
    # print(test)
    # print(test.text)
    # test.create_docx()

    # a = ReadPDF('./skany/test2')
    # a.read_pdf()
    # a.write_to_excel(a.handlers, 'test.xlsx')
    # for pdf in a.handlers:
    #     print(pdf.results['address'])
        # pdf.create_docx()

    parser = argparse.ArgumentParser(
        description="Script that generate administrarive decisions based on pdf files"
    )
    parser.add_argument("--path", required=True, type=str, help="Enter path to folder with pdf files")
    parser.add_argument("--scan", required=False, type=bool, default=False, help="Is the file a scan (image)")
    parser.add_argument("--reverse", required=False, type=bool, default=False, help="Write to excel in reverse order")

    args = parser.parse_args()

    a = ReadPDF(path=args.path, scan=args.scan, reverse=args.reverse)
    a.read_pdf()
    a.write_to_excel(a.handlers, 'test.xlsx')
    for pdf in a.handlers:
        pdf.create_docx()
